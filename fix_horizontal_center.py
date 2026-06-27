from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm

doc = Document("template.docx")
table5 = doc.tables[5]

# Center all paragraphs horizontally in table 5
for row in table5.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save("template.docx")
print("Horizontally centered texts in Table 5")
