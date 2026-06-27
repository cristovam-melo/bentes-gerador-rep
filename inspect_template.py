from docx import Document
import sys

try:
    doc = Document("template.docx")
    for para in doc.paragraphs:
        print(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            print(" | ".join(row_text))
except Exception as e:
    print("Error:", e)
