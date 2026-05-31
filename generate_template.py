from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx
import os

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_paragraph_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_run_font(run, name='Tahoma', size=10, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold

def set_cell_nowrap(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    noWrap = OxmlElement('w:noWrap')
    tcPr.append(noWrap)

def set_table_width_pct(table, pct=100.0, indent_cm=0):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), str(int(pct * 50)))
    
    if indent_cm > 0:
        tblInd = tblPr.first_child_found_in('w:tblInd')
        if tblInd is None:
            tblInd = OxmlElement('w:tblInd')
            tblPr.append(tblInd)
        tblInd.set(qn('w:type'), 'dxa')
        tblInd.set(qn('w:w'), str(int(indent_cm * 567)))

def set_cell_valign(cell, align='center'):
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = tcPr.first_child_found_in('w:vAlign')
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        tcPr.append(vAlign)
    vAlign.set(qn('w:val'), align)

doc = Document()

# Page setup: A4 Landscape, 2cm margins
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.header_distance = Cm(1.0)
section.footer_distance = Cm(1.0)

# Define Normal style to Tahoma
style = doc.styles['Normal']
font = style.font
font.name = 'Tahoma'
font.size = Pt(10)

# Logo insertion
p_logo = doc.add_paragraph()
if os.path.exists('logo.png'):
    run_logo = p_logo.add_run()
    run_logo.add_picture('logo.png', width=Inches(2.5))
elif os.path.exists('logo.jpg'):
    run_logo = p_logo.add_run()
    run_logo.add_picture('logo.jpg', width=Inches(2.5))
else:
    run_logo = p_logo.add_run("[ COLE O ARQUIVO logo.png OU logo.jpg AQUI NA PASTA E GERE O TEMPLATE NOVAMENTE ]\n")
    set_run_font(run_logo, size=12, bold=True)

set_paragraph_bottom_border(p_logo)

# Title and REP Number (using a table for alignment)
table_title = doc.add_table(rows=1, cols=2)
set_table_width_pct(table_title)
table_title.autofit = False
table_title.columns[0].width = Inches(4.0)
table_title.columns[1].width = Inches(2.0)

cell_title = table_title.cell(0, 0)
p_title = cell_title.paragraphs[0]
run_title = p_title.add_run("REGISTRO DE ENVIO DE PROJETOS")
set_run_font(run_title, size=16, bold=True)

cell_rep = table_title.cell(0, 1)
p_rep = cell_rep.paragraphs[0]
p_rep.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_rep = p_rep.add_run("REP- {{ numero_rep }}")
set_run_font(run_rep, size=16, bold=True)

doc.add_paragraph("") # Spacing

# Print 1: Info block -> Tahoma 10, Bold
table_info = doc.add_table(rows=2, cols=4)
set_table_width_pct(table_info)
data_info = [
    ("Data:", " {{ data_atual }}", "De:", " {{ remetente }}"),
    ("Cliente:", " {{ cliente }}", "Obra:", " {{ obra }}")
]
for r_idx, row_data in enumerate(data_info):
    for c_idx, text in enumerate(row_data):
        cell = table_info.cell(r_idx, c_idx)
        cell.text = ""
        run = cell.paragraphs[0].add_run(text)
        set_run_font(run, size=10, bold=True)

doc.add_paragraph("") # Spacing

# Recipients block
table_dest = doc.add_table(rows=4, cols=2)
set_table_width_pct(table_dest)
table_dest.style = 'Table Grid'

# Print 2: Recipients Headers -> Tahoma 9, Bold
hdr_cells = table_dest.rows[0].cells
hdr_cells[0].text = ""
p = hdr_cells[0].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Destinatário')
set_run_font(run, size=9, bold=True)

hdr_cells[1].text = ""
p = hdr_cells[1].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Empresa')
set_run_font(run, size=9, bold=True)

# Print 3: Recipients Data -> Tahoma 9, Normal
table_dest.rows[1].cells[0].text = ""
run = table_dest.rows[1].cells[0].paragraphs[0].add_run('{%tr for dest in destinatarios %}')

data_row_dest = table_dest.rows[2]
data_row_dest.cells[0].text = ""
p = data_row_dest.cells[0].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('{{ dest.nome }}')
set_run_font(run, size=9, bold=False)

data_row_dest.cells[1].text = ""
p = data_row_dest.cells[1].paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('{{ dest.empresa }}')
set_run_font(run, size=9, bold=False)

table_dest.rows[3].cells[0].text = ""
run = table_dest.rows[3].cells[0].paragraphs[0].add_run('{%tr endfor %}')

# Print 4: Forma do Envio
p_envio = doc.add_paragraph()
run_envio1 = p_envio.add_run("Forma do Envio/Tipo de Arquivo/Mídia: ")
set_run_font(run_envio1, size=9, bold=True)
run_envio2 = p_envio.add_run("Email/arquivo .pdf")
set_run_font(run_envio2, size=9, bold=False)

# Files block
table_files = doc.add_table(rows=4, cols=4)
set_table_width_pct(table_files, pct=97.2, indent_cm=0.04)
table_files.style = 'Table Grid'

# Print 5: Files Headers -> Tahoma 10, Bold, with line break on "Nº Folha" and thick borders
headers = ['Arquivo', 'Nº Folha', 'Descrição Folha', 'Data Folha']
hdr_files = table_files.rows[0].cells
for i, h_text in enumerate(headers):
    hdr_files[i].text = ""
    p = hdr_files[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h_text)
    set_run_font(run, size=10, bold=True)
    
    # Add thick border (sz=18 represents 2.25pt) to top and bottom of header row
    set_cell_border(hdr_files[i], 
                    top={"sz": 18, "val": "single"}, 
                    bottom={"sz": 18, "val": "single"})

# Print 6: Files Data -> Tahoma 10, Normal
table_files.rows[1].cells[0].text = ""
run = table_files.rows[1].cells[0].paragraphs[0].add_run('{%tr for file in arquivos %}')

data_row_files = table_files.rows[2]

# Apply noWrap to all columns except 'Descrição Folha' (index 2)
set_cell_nowrap(data_row_files.cells[0])
set_cell_nowrap(data_row_files.cells[1])
set_cell_nowrap(data_row_files.cells[3])
# Apply noWrap to the header cells as well so they don't wrap either
set_cell_nowrap(hdr_files[0])
set_cell_nowrap(hdr_files[1])
set_cell_nowrap(hdr_files[3])

data_fields = [
    '{{ file.arquivo }}', 
    '{{ file.numero_folha }}', 
    '{{ file.descricao }}', 
    '{{ file.data_folha }}'
]

for i, field in enumerate(data_fields):
    data_row_files.cells[i].text = ""
    p = data_row_files.cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(field)
    set_run_font(run, size=10, bold=False)

table_files.rows[3].cells[0].text = ""
run = table_files.rows[3].cells[0].paragraphs[0].add_run('{%tr endfor %}')

table_files.autofit = False

# Apply explicit widths and vertical center to all cells
for row in table_files.rows:
    row.cells[0].width = Cm(4.18)
    row.cells[1].width = Cm(1.5)
    row.cells[2].width = Cm(17.5)
    row.cells[3].width = Cm(2.5)
    for cell in row.cells:
        set_cell_valign(cell, 'center')

doc.save('template.docx')
print("template.docx gerado com sucesso!")
