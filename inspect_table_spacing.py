from docx import Document

doc = Document("template.docx")
table5 = doc.tables[5]

for i, row in enumerate(table5.rows[2:4]): # Check the dynamic row and maybe header
    print(f"Row {i+2}")
    for j, cell in enumerate(row.cells):
        for k, p in enumerate(cell.paragraphs):
            pf = p.paragraph_format
            print(f"  Cell {j} Para {k}: text='{p.text}', space_before={pf.space_before}, space_after={pf.space_after}, line_spacing={pf.line_spacing}, align={pf.alignment}")
