from docx import Document

doc = Document("template.docx")
cell0 = doc.tables[0].rows[0].cells[0]
cell1 = doc.tables[0].rows[0].cells[1]

# Modify cell 0 run
for p in cell0.paragraphs:
    for r in p.runs:
        if "REP-" in r.text:
            r.text = r.text.replace("REP-", "").strip()

# Modify cell 1 run
for p in cell1.paragraphs:
    for r in p.runs:
        if "{{ numero_rep }}" in r.text and "REP-" not in r.text:
            r.text = r.text.replace("{{ numero_rep }}", "REP- {{ numero_rep }}")

doc.save("template.docx")
print("Moved 'REP-' securely while keeping formatting.")
