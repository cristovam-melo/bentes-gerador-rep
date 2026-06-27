import docx
from docx import Document
from docxtpl import DocxTemplate

def main():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    
    p0 = row.cells[0].paragraphs[0]
    p0.add_run('{%tr for dest in destinatarios %}')
    p0_2 = row.cells[0].add_paragraph()
    p0_2.add_run('{{ dest.nome }}')
    
    p1 = row.cells[1].paragraphs[0]
    p1.add_run('{{ dest.empresa }}')
    p1_2 = row.cells[1].add_paragraph()
    p1_2.add_run('{%tr endfor %}')
    
    doc.save('test_in6.docx')

    tpl = DocxTemplate('test_in6.docx')
    context = {"destinatarios": [{"nome": "A", "empresa": "B"}]}
    try:
        tpl.render(context)
        tpl.save('test_out6.docx')
        print("SUCCESS 6")
    except Exception as e:
        print("FAILED 6:", e)

if __name__ == '__main__':
    main()
