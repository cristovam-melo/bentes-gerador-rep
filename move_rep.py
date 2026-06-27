from docx import Document

doc = Document("template.docx")
cell0 = doc.tables[0].rows[0].cells[0]
cell1 = doc.tables[0].rows[0].cells[1]

# Replace text in cell 0
for p in cell0.paragraphs:
    if "REP-" in p.text:
        p.text = p.text.replace("REP-", "").strip()

# Replace text in cell 1
for p in cell1.paragraphs:
    if "{{ numero_rep }}" in p.text and "REP-" not in p.text:
        p.text = p.text.replace("{{ numero_rep }}", "REP- {{ numero_rep }}")

doc.save("template.docx")
print("Moved 'REP-' from cell 0 to cell 1.")
