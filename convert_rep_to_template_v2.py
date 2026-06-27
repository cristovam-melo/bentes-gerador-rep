import docx
import copy
import sys

doc = docx.Document('temp.docx')

def replace_cell_text(cell, new_text):
    if cell.paragraphs and cell.paragraphs[0].runs:
        font_name = cell.paragraphs[0].runs[0].font.name
        font_size = cell.paragraphs[0].runs[0].font.size
        bold = cell.paragraphs[0].runs[0].bold
        color = cell.paragraphs[0].runs[0].font.color.rgb if cell.paragraphs[0].runs[0].font.color else None
    else:
        font_name, font_size, bold, color = None, None, None, None

    cell.text = new_text
    
    if cell.paragraphs and font_name:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = font_name
                if font_size: r.font.size = font_size
                r.bold = bold
                if color: r.font.color.rgb = color

# Table 0: REP number
replace_cell_text(doc.tables[0].cell(0, 1), '{{ numero_rep }}')

# Table 1: Data, Remetente
replace_cell_text(doc.tables[1].cell(0, 1), '{{ data_atual }}')
replace_cell_text(doc.tables[1].cell(0, 3), '{{ remetente }}')

# Table 2: Cliente, Obra
replace_cell_text(doc.tables[2].cell(0, 1), '{{ cliente }}')
replace_cell_text(doc.tables[2].cell(0, 3), '{{ obra }}')

# Table 3: Destinatarios
t3 = doc.tables[3]
for r in t3.rows[2:]:
    t3._tbl.remove(r._tr)

data_row_dest = t3.rows[1]
replace_cell_text(data_row_dest.cells[0], '{{ dest.nome }}')
replace_cell_text(data_row_dest.cells[2], '{{ dest.empresa }}')

# Deepcopy data_row_dest to preserve its EXACT NO-BORDER formatting
tr_for = copy.deepcopy(data_row_dest._tr)
tr_endfor = copy.deepcopy(data_row_dest._tr)

# Clear text and add tags
for tc in tr_for.findall('.//w:tc', namespaces=tr_for.nsmap):
    for p in tc.findall('.//w:p', namespaces=tc.nsmap):
        for r in p.findall('.//w:r', namespaces=p.nsmap):
            p.remove(r)
# Add text to first cell
p_for = tr_for.find('.//w:p', namespaces=tr_for.nsmap)
r_for = docx.oxml.OxmlElement('w:r')
t_for = docx.oxml.OxmlElement('w:t')
t_for.text = '{%tr for dest in destinatarios %}'
r_for.append(t_for)
p_for.append(r_for)

for tc in tr_endfor.findall('.//w:tc', namespaces=tr_endfor.nsmap):
    for p in tc.findall('.//w:p', namespaces=tc.nsmap):
        for r in p.findall('.//w:r', namespaces=p.nsmap):
            p.remove(r)
p_endfor = tr_endfor.find('.//w:p', namespaces=tr_endfor.nsmap)
r_endfor = docx.oxml.OxmlElement('w:r')
t_endfor = docx.oxml.OxmlElement('w:t')
t_endfor.text = '{%tr endfor %}'
r_endfor.append(t_endfor)
p_endfor.append(r_endfor)

t3._tbl.insert(t3._tbl.index(data_row_dest._tr), tr_for)
t3._tbl.append(tr_endfor)

# Table 5: Arquivos
t5 = doc.tables[5]
for r in t5.rows[2:]:
    t5._tbl.remove(r._tr)

data_row_file = t5.rows[1]
replace_cell_text(data_row_file.cells[0], '{{ file.arquivo }}')
replace_cell_text(data_row_file.cells[1], '{{ file.numero_folha }}')
replace_cell_text(data_row_file.cells[2], '{{ file.descricao }}')
replace_cell_text(data_row_file.cells[3], '{{ file.data_folha }}')

tr_for_file = copy.deepcopy(data_row_file._tr)
tr_endfor_file = copy.deepcopy(data_row_file._tr)

for tc in tr_for_file.findall('.//w:tc', namespaces=tr_for_file.nsmap):
    for p in tc.findall('.//w:p', namespaces=tc.nsmap):
        for r in p.findall('.//w:r', namespaces=p.nsmap):
            p.remove(r)
p_for_file = tr_for_file.find('.//w:p', namespaces=tr_for_file.nsmap)
r_for_file = docx.oxml.OxmlElement('w:r')
t_for_file = docx.oxml.OxmlElement('w:t')
t_for_file.text = '{%tr for file in arquivos %}'
r_for_file.append(t_for_file)
p_for_file.append(r_for_file)

for tc in tr_endfor_file.findall('.//w:tc', namespaces=tr_endfor_file.nsmap):
    for p in tc.findall('.//w:p', namespaces=tc.nsmap):
        for r in p.findall('.//w:r', namespaces=p.nsmap):
            p.remove(r)
p_endfor_file = tr_endfor_file.find('.//w:p', namespaces=tr_endfor_file.nsmap)
r_endfor_file = docx.oxml.OxmlElement('w:r')
t_endfor_file = docx.oxml.OxmlElement('w:t')
t_endfor_file.text = '{%tr endfor %}'
r_endfor_file.append(t_endfor_file)
p_endfor_file.append(r_endfor_file)

t5._tbl.insert(t5._tbl.index(data_row_file._tr), tr_for_file)
t5._tbl.append(tr_endfor_file)

doc.save('template_perfect_v2.docx')
print("Perfect template v2 generated!")
