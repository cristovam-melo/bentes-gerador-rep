from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

doc = Document("template.docx")

# Fix Table 0 (REP header)
# We want Col 0 to be wide and Col 1 to be wide enough so "REP- 001" doesn't wrap.
table0 = doc.tables[0]
table0.autofit = False
# In Word, if table is 100% width, we can set gridCol w or just cell widths.
for row in table0.rows:
    row.cells[0].width = Cm(13)
    row.cells[1].width = Cm(4)
# Make sure REP- 001 is right aligned (or left aligned if it's already in the right spot, but the original left image shows it right-aligned within its space or just having enough space)
for p in table0.rows[0].cells[1].paragraphs:
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Fix Table 5 (Files table)
table5 = doc.tables[5]
table5.autofit = False
for row in table5.rows:
    # Set explicit widths to prevent weird auto-scaling
    row.cells[0].width = Cm(5.5)  # Arquivo
    row.cells[1].width = Cm(1.5)  # Nº Folha
    row.cells[2].width = Cm(8.0)  # Descrição
    row.cells[3].width = Cm(2.0)  # Data Folha

# Revert Arquivo (Col 0) to left alignment and keep others centered
for i, row in enumerate(table5.rows):
    if i == 0:
        continue # skip header if needed, or adjust it too. Actually header "Arquivo" in original is centered!
    
    # Col 0: Left align, add left indent
    for p in row.cells[0].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.2)
        
    # Col 1, 2, 3: Center align
    for p in row.cells[1].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for p in row.cells[2].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for p in row.cells[3].paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save("template.docx")
print("Adjusted table widths and alignments.")
