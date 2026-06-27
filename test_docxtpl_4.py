import docx
from docx import Document
from docxtpl import DocxTemplate

def main():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    p1 = row.cells[0].paragraphs[0]
    p1.add_run('{%tr for dest in destinatarios %}')
    p2 = row.cells[0].add_paragraph()
    p2.add_run('{{ dest.nome }}')
    p3 = row.cells[0].add_paragraph()
    p3.add_run('{%tr endfor %}')
    
    row.cells[1].text = '{{ dest.empresa }}'
    doc.save('test_in4.docx')

    tpl = DocxTemplate('test_in4.docx')
    context = {"destinatarios": [{"nome": "A", "empresa": "B"}]}
    try:
        tpl.render(context)
        tpl.save('test_out4.docx')
        print("SUCCESS 4")
    except Exception as e:
        print("FAILED 4:", e)

if __name__ == '__main__':
    main()
