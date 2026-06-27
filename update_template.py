from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Cm, Inches

doc = Document("template.docx")

# 1. Blue rectangle: Table 0, Row 0 - Increase height and center vertically
table0 = doc.tables[0]
row0 = table0.rows[0]
if row0.height:
    row0.height = row0.height + Pt(10)
else:
    row0.height = Cm(1.5)
for cell in row0.cells:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# 2. Green rectangle: Table 3 (Destinatarios) - Align text to bottom
table3 = doc.tables[3]
# Row 2 contains {{ dest.nome }} and {{ dest.empresa }}
for cell in table3.rows[2].cells:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

# 3. Orange rectangle: Table 5 (Files) - Center vertically and widen column 0
table5 = doc.tables[5]
for row in table5.rows:
    for cell in row.cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

for row in table5.rows:
    c0 = row.cells[0]
    if c0.width:
        c0.width = c0.width + Cm(0.5)

for row in table5.rows[1:]: # Skip header
    for p in row.cells[0].paragraphs:
        p.paragraph_format.left_indent = Cm(0.2)

doc.save("template.docx")
print("Template updated successfully.")
