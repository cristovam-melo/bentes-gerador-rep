import docx
from docx import Document
from docxtpl import DocxTemplate

def main():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    row.cells[0].text = '{%tr for dest in destinatarios %}{{ dest.nome }}'
    row.cells[1].text = '{{ dest.empresa }}{%tr endfor %}'
    doc.save('test_in.docx')

    tpl = DocxTemplate('test_in.docx')
    context = {"destinatarios": [{"nome": "A", "empresa": "B"}]}
    try:
        tpl.render(context)
        tpl.save('test_out.docx')
        print("SUCCESS 1")
    except Exception as e:
        print("FAILED 1:", e)

    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    # Use separate paragraphs!
    p0_1 = row.cells[0].paragraphs[0]
    p0_1.add_run('{%tr for dest in destinatarios %}')
    p0_2 = row.cells[0].add_paragraph()
    p0_2.add_run('{{ dest.nome }}')
    
    p1_1 = row.cells[1].paragraphs[0]
    p1_1.add_run('{{ dest.empresa }}')
    p1_2 = row.cells[1].add_paragraph()
    p1_2.add_run('{%tr endfor %}')
    doc.save('test_in2.docx')

    tpl2 = DocxTemplate('test_in2.docx')
    try:
        tpl2.render(context)
        tpl2.save('test_out2.docx')
        print("SUCCESS 2")
    except Exception as e:
        print("FAILED 2:", e)

if __name__ == '__main__':
    main()
