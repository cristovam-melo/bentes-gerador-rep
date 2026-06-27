from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, name='Tahoma', size=10, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold

def set_table_width_pct(table, pct=100.0):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'pct')
    tblW.set(qn('w:w'), str(int(pct * 50)))

doc = Document('template_repo.docx')

# Find the destinatarios table
dest_table = None
for table in doc.tables:
    if len(table.rows) > 0 and "Destinat" in table.cell(0,0).text:
        dest_table = table
        break

if dest_table:
    # Clear all existing rows from XML
    tbl = dest_table._tbl
    for tr in list(tbl.xpath('.//w:tr')):
        tbl.remove(tr)
        
    # Rebuild the table rows and cells properly
    for i in range(4):
        dest_table.add_row()
        
    set_table_width_pct(dest_table)
    
    # Row 0: Headers
    r0c0 = dest_table.rows[0].cells[0]
    r0c0.text = ""
    p = r0c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Destinatário')
    set_run_font(run, size=9, bold=True)
    
    r0c1 = dest_table.rows[0].cells[1]
    r0c1.text = ""
    p = r0c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Empresa')
    set_run_font(run, size=9, bold=True)
    
    # Row 1: Loop start
    r1c0 = dest_table.rows[1].cells[0]
    r1c0.text = ""
    run = r1c0.paragraphs[0].add_run('{%tr for dest in destinatarios %}')
    
    # Row 2: Data
    r2c0 = dest_table.rows[2].cells[0]
    r2c0.text = ""
    p = r2c0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('{{ dest.nome }}')
    set_run_font(run, size=9, bold=False)
    
    r2c1 = dest_table.rows[2].cells[1]
    r2c1.text = ""
    p = r2c1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('{{ dest.empresa }}')
    set_run_font(run, size=9, bold=False)
    
    # Row 3: Loop end
    r3c0 = dest_table.rows[3].cells[0]
    r3c0.text = ""
    run = r3c0.paragraphs[0].add_run('{%tr endfor %}')

# We also need to remove the floating {%tr endfor %} paragraph
for p in doc.paragraphs:
    if "{%tr endfor %}" in p.text:
        # clear it
        p.text = p.text.replace("{%tr endfor %}", "")

doc.save('template_fixed.docx')
print("Fixed template saved.")
