import docx
from docx import Document
from docxtpl import DocxTemplate

def main():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    row = table.rows[0]
    
    p0 = row.cells[0].paragraphs[0]
    p0.add_run('{%tr for dest in destinatarios %}')
    p0_2 = row.cells[0].add_paragraph()
    p0_2.add_run('{{ dest.nome }}')
    p0_3 = row.cells[0].add_paragraph()
    p0_3.add_run('{%tr endfor %}')
    
    p1 = row.cells[1].paragraphs[0]
    p1.add_run('{{ dest.empresa }}')
    
    doc.save('test_in7.docx')

    tpl = DocxTemplate('test_in7.docx')
    context = {"destinatarios": [{"nome": "A", "empresa": "B"}]}
    try:
        tpl.render(context)
        tpl.save('test_out7.docx')
        print("SUCCESS 7")
    except Exception as e:
        print("FAILED 7:", e)

if __name__ == '__main__':
    main()
