import docx
import sys

doc = docx.Document('REP-144.docx')

def replace_cell_text(cell, new_text):
    if cell.paragraphs and cell.paragraphs[0].runs:
        font_name = cell.paragraphs[0].runs[0].font.name
        font_size = cell.paragraphs[0].runs[0].font.size
        bold = cell.paragraphs[0].runs[0].bold
        color = cell.paragraphs[0].runs[0].font.color.rgb if cell.paragraphs[0].runs[0].font.color else None
    else:
        font_name, font_size, bold, color = None, None, None, None

    cell.text = new_text
    
    if cell.paragraphs and font_name:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = font_name
                if font_size: r.font.size = font_size
                r.bold = bold
                if color: r.font.color.rgb = color

# Table 0: REP number
replace_cell_text(doc.tables[0].cell(0, 1), '{{ numero_rep }}')

# Table 1: Data, Remetente
replace_cell_text(doc.tables[1].cell(0, 1), '{{ data_atual }}')
replace_cell_text(doc.tables[1].cell(0, 3), '{{ remetente }}')

# Table 2: Cliente, Obra
replace_cell_text(doc.tables[2].cell(0, 1), '{{ cliente }}')
replace_cell_text(doc.tables[2].cell(0, 3), '{{ obra }}')

# Table 3: Destinatarios
t3 = doc.tables[3]
# Preserve row 1 as the data row, delete any rows after row 1
rows_to_delete = t3.rows[2:]
for r in rows_to_delete:
    t3._tbl.remove(r._tr)

data_row_dest = t3.rows[1]
replace_cell_text(data_row_dest.cells[0], '{{ dest.nome }}')
replace_cell_text(data_row_dest.cells[2], '{{ dest.empresa }}')

row_for_dest = t3.add_row()
row_for_dest.cells[0].text = '{%tr for dest in destinatarios %}'

row_endfor_dest = t3.add_row()
row_endfor_dest.cells[0].text = '{%tr endfor %}'

# Move row_for_dest BEFORE data_row_dest
t3._tbl.insert(t3._tbl.index(data_row_dest._tr), row_for_dest._tr)

# Table 5: Arquivos
t5 = doc.tables[5]
# Preserve row 1 as the data row, delete any rows after row 1
rows_to_delete = t5.rows[2:]
for r in rows_to_delete:
    t5._tbl.remove(r._tr)

data_row_file = t5.rows[1]
replace_cell_text(data_row_file.cells[0], '{{ file.arquivo }}')
replace_cell_text(data_row_file.cells[1], '{{ file.numero_folha }}')
replace_cell_text(data_row_file.cells[2], '{{ file.descricao }}')
replace_cell_text(data_row_file.cells[3], '{{ file.data_folha }}')

row_for_file = t5.add_row()
row_for_file.cells[0].text = '{%tr for file in arquivos %}'

row_endfor_file = t5.add_row()
row_endfor_file.cells[0].text = '{%tr endfor %}'

# Move row_for_file BEFORE data_row_file
t5._tbl.insert(t5._tbl.index(data_row_file._tr), row_for_file._tr)

doc.save('template_perfect.docx')
print("Perfect template generated!")
