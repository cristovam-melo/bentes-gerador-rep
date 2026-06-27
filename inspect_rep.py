from docx import Document

doc = Document("template.docx")
table0 = doc.tables[0]
cell = table0.rows[0].cells[1]

print("Cell text:", repr(cell.text))
for i, p in enumerate(cell.paragraphs):
    print(f"Para {i} alignment: {p.alignment}")
    for j, r in enumerate(p.runs):
        print(f"  Run {j}: {repr(r.text)}")
